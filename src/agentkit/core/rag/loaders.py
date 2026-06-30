"""Filesystem document loading for enterprise RAG ingestion.

The loader is intentionally dependency-light at import time. PDF, Word, OCR and
image analysis dependencies are imported only when the matching file type or
option is used.
"""

from __future__ import annotations

import hashlib
import html
import io
import json
import re
import zipfile
from collections.abc import Iterable, Sequence
from dataclasses import dataclass, field
from pathlib import Path
from typing import Protocol

from .base import KnowledgeDocument

SUPPORTED_EXTENSIONS = {
    ".csv",
    ".docx",
    ".htm",
    ".html",
    ".json",
    ".md",
    ".pdf",
    ".txt",
}


class ImageAnalyzer(Protocol):
    def analyze(self, image_bytes: bytes, *, mime_type: str, hint: str = "") -> str: ...


@dataclass(frozen=True)
class DocumentLoadOptions:
    recursive: bool = True
    ocr_enabled: bool = False
    ocr_languages: str = "eng+chi_sim"
    min_page_text_chars: int = 40
    include_hidden_files: bool = False
    max_file_bytes: int = 50 * 1024 * 1024


@dataclass(frozen=True)
class FileLoadReport:
    documents: list[KnowledgeDocument] = field(default_factory=list)
    skipped: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


class TesseractImageAnalyzer:
    """OCR image bytes with pytesseract.

    Requires the Python packages from ``agentkit[rag]`` and a system tesseract
    binary. Docker images can install ``tesseract-ocr`` plus language packs.
    """

    def __init__(self, *, languages: str = "eng+chi_sim") -> None:
        self._languages = languages

    def analyze(self, image_bytes: bytes, *, mime_type: str, hint: str = "") -> str:
        try:
            import pytesseract
            from PIL import Image
        except ImportError as exc:  # pragma: no cover - optional dependency
            raise RuntimeError(
                "OCR requires the RAG optional dependencies. Install with: "
                "pip install 'agentkit[rag]'"
            ) from exc
        image = Image.open(io.BytesIO(image_bytes))
        text = pytesseract.image_to_string(image, lang=self._languages)
        return _clean_text(text)


class DocumentFolderLoader:
    """Load a file or folder into KnowledgeDocument objects.

    The returned documents include ``metadata["blocks"]`` for page/table/OCR
    aware chunking. Plain text remains a single document.
    """

    def __init__(
        self,
        *,
        options: DocumentLoadOptions | None = None,
        image_analyzer: ImageAnalyzer | None = None,
    ) -> None:
        self._options = options or DocumentLoadOptions()
        self._image_analyzer = image_analyzer
        if self._options.ocr_enabled and self._image_analyzer is None:
            self._image_analyzer = TesseractImageAnalyzer(languages=self._options.ocr_languages)

    def load_path(
        self,
        path: str | Path,
        *,
        tenant_id: str,
        acl_roles: Sequence[str] = (),
        metadata: dict | None = None,
    ) -> list[KnowledgeDocument]:
        return self.load_path_with_report(
            path,
            tenant_id=tenant_id,
            acl_roles=acl_roles,
            metadata=metadata,
        ).documents

    def load_path_with_report(
        self,
        path: str | Path,
        *,
        tenant_id: str,
        acl_roles: Sequence[str] = (),
        metadata: dict | None = None,
    ) -> FileLoadReport:
        root = Path(path)
        warnings: list[str] = []
        skipped: list[str] = []
        documents: list[KnowledgeDocument] = []
        for file_path in self._iter_files(root):
            try:
                documents.append(
                    self.load_file(
                        file_path,
                        tenant_id=tenant_id,
                        acl_roles=acl_roles,
                        metadata=metadata,
                    )
                )
            except ValueError as exc:
                skipped.append(f"{file_path}: {exc}")
            except RuntimeError as exc:
                warnings.append(f"{file_path}: {exc}")
        return FileLoadReport(documents=documents, skipped=skipped, warnings=warnings)

    def load_file(
        self,
        path: str | Path,
        *,
        tenant_id: str,
        acl_roles: Sequence[str] = (),
        metadata: dict | None = None,
    ) -> KnowledgeDocument:
        file_path = Path(path)
        if not file_path.is_file():
            raise ValueError("not a file")
        if file_path.stat().st_size > self._options.max_file_bytes:
            raise ValueError(f"file exceeds max_file_bytes={self._options.max_file_bytes}")
        suffix = file_path.suffix.lower()
        if suffix not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"unsupported extension: {suffix or '(none)'}")

        raw_bytes = file_path.read_bytes()
        file_hash = hashlib.sha256(raw_bytes).hexdigest()
        base_metadata = {
            "source_path": str(file_path),
            "source_name": file_path.name,
            "extension": suffix,
            "file_sha256": file_hash,
            **(metadata or {}),
        }
        if suffix == ".pdf":
            text, blocks, parser_warnings = self._extract_pdf(file_path)
        elif suffix == ".docx":
            text, blocks, parser_warnings = self._extract_docx(file_path)
        elif suffix in {".html", ".htm"}:
            text = _html_to_text(raw_bytes.decode("utf-8", errors="replace"))
            blocks = [{"text": text, "kind": "text", "source": "html"}] if text else []
            parser_warnings = []
        else:
            text = raw_bytes.decode("utf-8", errors="replace")
            text = _clean_text(text)
            blocks = [{"text": text, "kind": "text", "source": suffix.lstrip(".")}] if text else []
            parser_warnings = []

        doc_metadata = {
            **base_metadata,
            "blocks": blocks,
            "parser_warnings": parser_warnings,
        }
        return KnowledgeDocument(
            id=f"{tenant_id}:{file_hash[:24]}",
            tenant_id=tenant_id,
            text=text,
            title=file_path.stem,
            uri=file_path.resolve().as_uri(),
            metadata=doc_metadata,
            acl_roles=tuple(str(role) for role in acl_roles),
        )

    def _iter_files(self, root: Path) -> Iterable[Path]:
        if root.is_file():
            if self._include(root):
                yield root
            return
        if not root.is_dir():
            raise ValueError(f"path does not exist: {root}")
        iterator = root.rglob("*") if self._options.recursive else root.glob("*")
        for path in sorted(iterator):
            if path.is_file() and self._include(path):
                yield path

    def _include(self, path: Path) -> bool:
        if not self._options.include_hidden_files and any(
            part.startswith(".") for part in path.parts
        ):
            return False
        if path.name.startswith("~$"):
            return False
        return path.suffix.lower() in SUPPORTED_EXTENSIONS

    def _extract_pdf(self, path: Path) -> tuple[str, list[dict], list[str]]:
        warnings: list[str] = []
        blocks: list[dict] = []
        page_texts = self._extract_pdf_text(path, warnings)
        page_count = len(page_texts)
        if self._options.ocr_enabled:
            ocr_blocks = self._extract_pdf_ocr(path, page_texts, warnings)
            blocks.extend(ocr_blocks)
        for page_number, text in enumerate(page_texts, start=1):
            clean = _clean_text(text)
            if clean:
                blocks.append(
                    {
                        "text": clean,
                        "kind": "page_text",
                        "page": page_number,
                        "source": "pdf_text",
                    }
                )
        if page_count and not blocks:
            warnings.append("PDF contained no extractable text; enable OCR for scanned files.")
        blocks.sort(key=lambda block: (int(block.get("page") or 0), str(block.get("kind") or "")))
        return "\f".join(str(block["text"]) for block in blocks), blocks, warnings

    def _extract_pdf_text(self, path: Path, warnings: list[str]) -> list[str]:
        try:
            from pypdf import PdfReader
        except ImportError:
            return self._extract_pdf_text_with_fitz(path, warnings)
        try:
            reader = PdfReader(str(path))
            return [page.extract_text() or "" for page in reader.pages]
        except Exception as exc:  # noqa: BLE001
            warnings.append(f"pypdf extraction failed: {exc}")
            return self._extract_pdf_text_with_fitz(path, warnings)

    def _extract_pdf_text_with_fitz(self, path: Path, warnings: list[str]) -> list[str]:
        try:
            import fitz
        except ImportError as exc:  # pragma: no cover - optional dependency
            raise RuntimeError(
                "PDF ingestion requires pypdf or PyMuPDF. Install with: "
                "pip install 'agentkit[rag]'"
            ) from exc
        try:
            with fitz.open(path) as doc:
                return [page.get_text("text") or "" for page in doc]
        except Exception as exc:  # noqa: BLE001
            warnings.append(f"PyMuPDF text extraction failed: {exc}")
            return []

    def _extract_pdf_ocr(
        self,
        path: Path,
        page_texts: Sequence[str],
        warnings: list[str],
    ) -> list[dict]:
        if self._image_analyzer is None:
            return []
        try:
            import fitz
        except ImportError as exc:  # pragma: no cover - optional dependency
            raise RuntimeError(
                "PDF OCR requires PyMuPDF. Install with: pip install 'agentkit[rag]'"
            ) from exc
        out: list[dict] = []
        with fitz.open(path) as doc:
            for page_index, page in enumerate(doc, start=1):
                existing = page_texts[page_index - 1] if page_index <= len(page_texts) else ""
                if len(_clean_text(existing)) >= self._options.min_page_text_chars:
                    continue
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                image_bytes = pix.tobytes("png")
                try:
                    text = self._image_analyzer.analyze(
                        image_bytes,
                        mime_type="image/png",
                        hint=f"{path.name} page {page_index}",
                    )
                except Exception as exc:  # noqa: BLE001
                    warnings.append(f"OCR failed on page {page_index}: {exc}")
                    continue
                if text:
                    out.append(
                        {
                            "text": text,
                            "kind": "page_ocr",
                            "page": page_index,
                            "source": "ocr",
                        }
                    )
        return out

    def _extract_docx(self, path: Path) -> tuple[str, list[dict], list[str]]:
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover - optional dependency
            raise RuntimeError(
                "Word .docx ingestion requires python-docx. Install with: "
                "pip install 'agentkit[rag]'"
            ) from exc

        warnings: list[str] = []
        blocks: list[dict] = []
        document = Document(str(path))
        for paragraph in document.paragraphs:
            text = _clean_text(paragraph.text)
            if text:
                blocks.append({"text": text, "kind": "paragraph", "source": "python-docx"})
        for table_index, table in enumerate(document.tables, start=1):
            rows: list[str] = []
            for row in table.rows:
                cells = [_clean_text(cell.text) for cell in row.cells]
                rows.append(" | ".join(cell for cell in cells if cell))
            text = "\n".join(row for row in rows if row)
            if text:
                blocks.append(
                    {
                        "text": text,
                        "kind": "table",
                        "source": "python-docx",
                        "metadata": {"table_index": table_index},
                    }
                )
        if self._options.ocr_enabled and self._image_analyzer is not None:
            blocks.extend(self._extract_docx_image_ocr(path, warnings))
        text = "\n\n".join(str(block["text"]) for block in blocks)
        return text, blocks, warnings

    def _extract_docx_image_ocr(self, path: Path, warnings: list[str]) -> list[dict]:
        analyzer = self._image_analyzer
        if analyzer is None:
            return []
        out: list[dict] = []
        with zipfile.ZipFile(path) as archive:
            for name in sorted(archive.namelist()):
                if not name.startswith("word/media/"):
                    continue
                suffix = Path(name).suffix.lower().lstrip(".") or "png"
                mime_type = f"image/{'jpeg' if suffix in {'jpg', 'jpeg'} else suffix}"
                try:
                    text = analyzer.analyze(
                        archive.read(name),
                        mime_type=mime_type,
                        hint=f"{path.name}:{name}",
                    )
                except Exception as exc:  # noqa: BLE001
                    warnings.append(f"OCR failed on embedded image {name}: {exc}")
                    continue
                if text:
                    out.append(
                        {
                            "text": text,
                            "kind": "image_ocr",
                            "source": "ocr",
                            "metadata": {"image_name": name},
                        }
                    )
        return out


def _html_to_text(raw: str) -> str:
    raw = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", raw)
    raw = re.sub(r"(?s)<[^>]+>", " ", raw)
    return _clean_text(html.unescape(raw))


def _clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t\r\f\v]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def load_eval_dataset(path: str | Path) -> list[dict]:
    """Load JSON or JSONL files used by RAG evaluation CLI."""

    p = Path(path)
    raw = p.read_text(encoding="utf-8")
    if p.suffix.lower() == ".jsonl":
        return [json.loads(line) for line in raw.splitlines() if line.strip()]
    data = json.loads(raw)
    if isinstance(data, list):
        return [item for item in data if isinstance(item, dict)]
    if isinstance(data, dict) and isinstance(data.get("cases"), list):
        return [item for item in data["cases"] if isinstance(item, dict)]
    raise ValueError("RAG eval dataset must be a JSON list, {'cases': [...]}, or JSONL.")


__all__ = [
    "DocumentFolderLoader",
    "DocumentLoadOptions",
    "FileLoadReport",
    "ImageAnalyzer",
    "SUPPORTED_EXTENSIONS",
    "TesseractImageAnalyzer",
    "load_eval_dataset",
]
